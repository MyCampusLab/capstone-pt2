import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_capstone_report():
    doc = Document()
    
    # Configure 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Base Normal Style (Arial)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(38, 38, 38) # Charcoal dark gray

    # Helper function for cell background color
    def set_cell_shading(cell, color_hex):
        shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

    # Helper for cell padding (margins)
    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180): # in dxa
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("LAPORAN PROGRESS CAPSTONE PROJECT")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(27, 79, 114) # Deep Blue

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Sistem VisionSafe (EyeGuardian): Mitigasi Risiko Mata Minus Berbasis AI")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(127, 140, 141) # Muted Gray

    doc.add_paragraph() # Spacer

    # Collaborator details card style
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(4)
    run_meta_hdr = p_meta.add_run("IDENTITAS TIM PENGEMBANG:")
    run_meta_hdr.font.bold = True
    run_meta_hdr.font.size = Pt(11)
    run_meta_hdr.font.color.rgb = RGBColor(44, 62, 80)
    
    p_meta_body = doc.add_paragraph()
    p_meta_body.paragraph_format.line_spacing = 1.15
    run_irsyad = p_meta_body.add_run("1. M. Irsyad Fachryanto (NIM: 23090151)")
    run_irsyad.font.bold = True
    p_meta_body.add_run(" - Software Engineer & AI Core Developer\n   (Tanggung jawab: Algoritma AI Face Mesh, Android Native Background Service, Supabase Database & Cloud Triggers)\n")
    run_marsha = p_meta_body.add_run("2. Marsha Dwi Lucyana (NIM: 23090151)")
    run_marsha.font.bold = True
    p_meta_body.add_run(" - UI/UX Designer & Quality Assurance (QA)\n   (Tanggung jawab: Sederhana & Ringan - Perancangan Visual & Flowchart, Dokumentasi API Swagger/Postman, Unit/Widget Testing, Laporan Sidang)")

    doc.add_paragraph() # Spacer

    # Section Table Title
    p_sect = doc.add_paragraph()
    run_sect = p_sect.add_run("DISTRIBUSI TUGAS DAN PROGRES REALTIME (20 BARIS TUGAS)")
    run_sect.font.bold = True
    run_sect.font.size = Pt(12)
    run_sect.font.color.rgb = RGBColor(27, 79, 114)

    # 20 Tasks aligned with actual git modifications/creation timestamps (late April - June 2026)
    tasks = [
        # (No, Nama Task, Deskripsi, PIC, Prioritas, Status, Mulai, Deadline, Progress, Catatan)
        ("1", "Analisis Kelayakan AI & Pemilihan Model", "Riset performa model face_landmarker.task Google MediaPipe serta batas minimum FPS pada perangkat Android.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "20-04-2026", "25-04-2026", "100%", "Model terbukti ringan (<20ms per frame)"),
        ("2", "Perancangan Flowchart Sistem & Wireframe", "Menyusun flowchart alur deteksi latar belakang (foreground service) dan mockup visual awal aplikasi.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "22-04-2026", "26-04-2026", "100%", "Flowchart disetujui untuk panduan coding"),
        ("3", "Inisialisasi Project & Core UI Setup", "Melakukan setup awal project Flutter, manajemen routes, assets gambar/animasi, dan environment.", "Marsha Dwi Lucyana", "Tinggi", "Selesai", "27-04-2026", "02-05-2026", "100%", "Flutter project siap dengan struktur dasar"),
        ("4", "Setup Database Supabase", "Membuat tabel telemetry, users, dan konfigurasi PostgreSQL Row Level Security (RLS) di Cloud.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "29-04-2026", "03-05-2026", "100%", "RLS berjalan untuk proteksi data pengguna"),
        ("5", "Integrasi Supabase Auth", "Mengimplementasikan login dan registrasi email serta Google OAuth di Flutter.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "02-05-2026", "07-05-2026", "100%", "Autentikasi cloud berjalan lancar"),
        ("6", "Desain UI Dashboard Utama", "Mendesain layout dashboard utama, kartu ringkasan, dan widget statis dashboard.", "Marsha Dwi Lucyana", "Tinggi", "Selesai", "01-05-2026", "06-05-2026", "100%", "Desain dashboard rapi dengan visual modern"),
        ("7", "Integrasi MediaPipe Face Mesh", "Inisialisasi AI Core SDK Google MediaPipe dan penempatan file model face_landmarker.task.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "04-05-2026", "10-05-2026", "100%", "AI mendeteksi 478 titik koordinat wajah"),
        ("8", "Logika Estimasi Jarak", "Menerapkan rumus matematika Triangle Similarity (jarak 3D) dari koordinat mata ke layar.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "06-05-2026", "11-05-2026", "100%", "Perhitungan jarak teruji akurat"),
        ("9", "Implementasi Smoothing", "Menerapkan filter matematika (Low-Pass Filter) pada pembacaan jarak untuk meredam jitter.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "09-05-2026", "12-05-2026", "100%", "Nilai fluktuasi jarak berhasil dihaluskan"),
        ("10", "Mascot & Quest UI Styling", "Styling visual robot animasi maskot Vizo dan menyusun list quest/misi agar rapi di layar.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "10-05-2026", "13-05-2026", "100%", "Selesai pada commit b11f7f4"),
        ("11", "Setup Izin Android & Native Bridges", "Mengatur izin overlay window (draw over apps) dan inisialisasi MethodChannel Flutter-Kotlin.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "11-05-2026", "15-05-2026", "100%", "Aplikasi diizinkan berjalan di atas aplikasi lain"),
        ("12", "Handler updateThreshold di MainActivity", "Implementasi handler Kotlin native di MainActivity untuk update batas jarak secara dinamis.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "14-05-2026", "18-05-2026", "100%", "Selesai pada commit 09d74d8 & 5da3eba"),
        ("13", "Refactoring Sticker & Hive Migration", "Migrasi penyimpanan data lokal (quest/sticker) ke local database Hive untuk performa cepat.", "M. Irsyad Fachryanto", "Sedang", "Selesai", "13-05-2026", "17-05-2026", "100%", "Selesai pada commit 35ed55a"),
        ("14", "Integrasi Dynamic Ktor Backend URL", "Menambahkan modularitas Ktor client agar base URL backend API dapat diubah secara dinamis.", "M. Irsyad Fachryanto", "Sedang", "Selesai", "15-05-2026", "19-05-2026", "100%", "Selesai pada commit a35e24a"),
        ("15", "UI/UX Audit & Dependency Resolution", "Mengecek kecocokan font, warna, serta merapikan package dependencies pada pubspec.yaml.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "15-05-2026", "19-05-2026", "100%", "Selesai pada commit d683c50"),
        ("16", "Dokumentasi API Swagger UI", "Menyusun skema OpenAPI dan hosting dokumentasi interaktif di visionsafe-api.surge.sh.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "20-05-2026", "24-05-2026", "100%", "Dokumentasi aktif di visionsafe-api.surge.sh"),
        ("17", "Postman Collection Setup", "Membuat skenario pengujian API (Login, Register, Telemetry) untuk keperluan testing dosen.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "22-05-2026", "26-05-2026", "100%", "File JSON Postman diekspor ke folder api_docs"),
        ("18", "Skrip Automated API Tester", "Menyusun skrip sederhana dengan Python untuk hit endpoint API dan verifikasi respon server.", "Marsha Dwi Lucyana", "Sedang", "Selesai", "24-05-2026", "27-05-2026", "100%", "Script python tersimpan di scratch/test_api_supabase.py"),
        ("19", "Integrasi Sistem Intervensi Blur & Lock", "Menyempurnakan intervensi visual blur overlay dan penguncian emergency jika pelanggaran >10 detik.", "M. Irsyad Fachryanto", "Tinggi", "Selesai", "28-05-2026", "02-06-2026", "100%", "Berhasil diuji dengan system overlay Android"),
        ("20", "Laporan Capstone & PPT", "Menulis draft dokumen Bab 4-5 (Hasil/Uji coba) dan menyiapkan slide presentasi sidang.", "Marsha Dwi Lucyana", "Sedang", "On Progress", "02-06-2026", "15-06-2026", "75%", "Draft laporan dan PPT sedang diselesaikan")
    ]

    # Create Table: 21 rows (1 header + 20 tasks), 10 columns
    headers = ["No", "Nama Task", "Deskripsi", "PIC (Penanggung Jawab)", "Prioritas", "Status", "Tanggal Mulai", "Deadline", "Progress (%)", "Catatan"]
    col_widths = [Inches(0.4), Inches(1.3), Inches(1.8), Inches(1.1), Inches(0.7), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.7), Inches(1.2)]
    
    table = doc.add_table(rows=1, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Design Header Row
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_shading(hdr_cells[i], "1F4E79") # Professional Navy Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        
        # Style header font
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                run.font.size = Pt(8.5)

    # Populate Data Rows
    for row_idx, task_data in enumerate(tasks):
        row_cells = table.add_row().cells
        # Zebra striping background (alternate gray)
        bg_color = "F2F4F4" if row_idx % 2 == 1 else "FFFFFF"
        
        for col_idx, text_val in enumerate(task_data):
            row_cells[col_idx].text = text_val
            set_cell_shading(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            
            # Format text per cell
            for paragraph in row_cells[col_idx].paragraphs:
                # Text alignment options
                if col_idx in [0, 4, 5, 6, 7, 8]: # Centered columns (No, Priority, Status, Dates, Progress)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    # Color codes for status
                    if col_idx == 5: # Status Column
                        if text_val == "Selesai":
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(39, 174, 96) # Green
                        elif text_val == "On Progress":
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(230, 126, 34) # Orange
                    elif col_idx == 3: # PIC Column
                        if "Irsyad" in text_val:
                            run.font.color.rgb = RGBColor(41, 128, 185) # Blue
                        else:
                            run.font.color.rgb = RGBColor(142, 68, 173) # Purple
                            
    # Set explicit column widths for all cells in all rows
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Save output to workspace
    output_path = "/home/irsyad/Gudang/EyeGuardian/visionsafe/Laporan_Progress_Capstone_VisionSafe.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX file at: {output_path}")

if __name__ == "__main__":
    create_capstone_report()
