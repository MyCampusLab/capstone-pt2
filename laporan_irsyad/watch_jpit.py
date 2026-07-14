import os
import re
import time
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def configure_styles(doc):
    # Base Normal Style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1)

    # Heading 1 (BAB)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(10)
    h1.font.bold = True
    h1.font.small_caps = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)

    # Heading 2 (SUB BAB)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(10)
    h2.font.italic = True
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(0)
    
    # Code Block Style
    if 'Code' not in doc.styles:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles['Code']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.first_line_indent = Cm(0)
    code_style.paragraph_format.left_indent = Cm(1)

def add_paragraph_with_inline_formatting(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    if style == 'Normal':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    # Regex to find *italic* or `backtick`
    # Both are transformed into italic text, removing the symbols.
    tokens = re.split(r'(\*[^*]+\*|`[^`]+`)', text)
    for token in tokens:
        if (token.startswith('*') and token.endswith('*')) or (token.startswith('`') and token.endswith('`')):
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)

class JournalCompiler:
    def __init__(self):
        self.roman_counter = 1
        self.roman_numerals = {1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V'}

    def parse_markdown_to_docx(self, md_path, doc):
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Ignore guides separated by ---
        content = content.split('---')[0].strip()
        
        lines = content.split('\n')
        in_code_block = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                doc.add_paragraph(line, style='Code')
                continue
                
            if line.startswith('`') and line.endswith('`') and len(line) > 2:
                # Menangani baris penuh kode yang dibungkus dengan single backtick
                doc.add_paragraph(line[1:-1], style='Code')
                continue
                
            if line.startswith('## '):
                add_paragraph_with_inline_formatting(doc, line[3:], style='Heading 2')
            elif line.startswith('# '):
                header_text = line[2:].strip().upper()
                if header_text in ["PENDAHULUAN", "METODE", "HASIL"]:
                    formatted_text = f"{self.roman_numerals.get(self.roman_counter, '?')}. {header_text}"
                    self.roman_counter += 1
                else:
                    formatted_text = header_text # e.g. SIMPULAN
                add_paragraph_with_inline_formatting(doc, formatted_text, style='Heading 1')
            else:
                # Need to handle specific formatting like (Tempatkan Gambar 1 Di Sini) 
                # Let's just output as normal paragraph for now
                add_paragraph_with_inline_formatting(doc, line, style='Normal')

def compile_journal():
    doc = Document()
    configure_styles(doc)
    
    files_to_compile = [
        'irsyad_04_B_Bab_Pendahuluan.md',
        'irsyad_05_Bab_Metode.md',
        'irsyad_06_Bab_Hasil.md',
        'irsyad_07_Bab_Simpulan.md'
    ]
    
    compiler = JournalCompiler()
    
    for f in files_to_compile:
        if os.path.exists(f):
            compiler.parse_markdown_to_docx(f, doc)
            doc.add_paragraph('') # Empty paragraph spacing between files
    
    doc.save('AUTO_BUILD_JPIT_2024.docx')
    print("Compiled successfully to AUTO_BUILD_JPIT_2024.docx")

if __name__ == '__main__':
    print("Memulai Mesin Otomatisasi JPIT 2024 (Watching for changes...)")
    print("Tekan Ctrl+C untuk menghentikan.")
    last_mtimes = {}
    
    # Initial compile
    compile_journal()
    for f in os.listdir('.'):
        if f.endswith('.md') and f.startswith('irsyad_'):
            last_mtimes[f] = os.path.getmtime(f)
            
    while True:
        changed = False
        for f in os.listdir('.'):
            if f.endswith('.md') and f.startswith('irsyad_'):
                mtime = os.path.getmtime(f)
                if f not in last_mtimes or mtime > last_mtimes[f]:
                    last_mtimes[f] = mtime
                    changed = True
        
        if changed:
            print("Perubahan terdeteksi! Membangun ulang jurnal ke JPIT 2024...")
            try:
                compile_journal()
            except Exception as e:
                print(f"Gagal membangun: {e}")
                
        time.sleep(2)
